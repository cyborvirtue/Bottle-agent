#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文档处理模块
支持PDF、TXT、Markdown等格式的文档解析和分块
"""

import os
import hashlib
from pathlib import Path
from typing import List, Dict, Any, Optional
import logging
from dataclasses import dataclass

# PDF处理
try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("⚠️  PDF处理库未安装，将跳过PDF文件")

# Word文档处理
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️  Word文档处理库未安装，将跳过DOCX文件")

from .data_structures import DocumentChunk, Document


class DocumentProcessor:
    """文档处理器"""
    
    def __init__(self, config: Dict[str, Any]):
        self.config = config
        self.rag_config = config["rag"]
        self.kb_config = config["knowledge_base"]
        
        # 支持的文件格式
        self.supported_formats = self.kb_config["supported_formats"]
        
        # 分块参数
        self.chunk_size = self.rag_config["chunk_size"]
        self.chunk_overlap = self.rag_config["chunk_overlap"]
    
    def process_folder(self, folder_path: str) -> List[Document]:
        """处理文件夹中的所有文档
        
        Args:
            folder_path: 文件夹路径
            
        Returns:
            文档列表
        """
        folder_path = Path(folder_path)
        documents = []
        
        if not folder_path.exists():
            logging.error(f"文件夹不存在: {folder_path}")
            return documents
        
        # 遍历文件夹
        for file_path in folder_path.rglob("*"):
            if file_path.is_file() and file_path.suffix.lower() in self.supported_formats:
                try:
                    document = self.process_file(str(file_path))
                    if document:
                        documents.append(document)
                        print(f"✅ 处理文件: {file_path.name}")
                except Exception as e:
                    logging.error(f"处理文件失败 {file_path}: {e}")
                    print(f"❌ 处理文件失败: {file_path.name} - {e}")
        
        return documents
    
    def process_file(self, file_path: str) -> Optional[Document]:
        """处理单个文件
        
        Args:
            file_path: 文件路径
            
        Returns:
            文档对象
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            logging.error(f"文件不存在: {file_path}")
            return None
        
        # 根据文件扩展名选择处理方法
        suffix = file_path.suffix.lower()
        
        if suffix == ".pdf":
            return self._process_pdf(file_path)
        elif suffix == ".txt":
            return self._process_txt(file_path)
        elif suffix == ".md":
            return self._process_markdown(file_path)
        elif suffix == ".docx":
            return self._process_docx(file_path)
        else:
            logging.warning(f"不支持的文件格式: {suffix}")
            return None
    
    def _process_pdf(self, file_path: Path) -> Optional[Document]:
        """处理PDF文件
        
        Args:
            file_path: PDF文件路径
            
        Returns:
            文档对象
        """
        if not PDF_AVAILABLE:
            logging.warning("PDF处理库未安装")
            return None
        
        content = ""
        metadata = {
            "source": file_path.name,
            "file_path": str(file_path),
            "file_type": "pdf",
            "file_size": file_path.stat().st_size
        }
        
        try:
            # 尝试使用pdfplumber（更好的文本提取）
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        content += f"\n\n[第{page_num}页]\n{page_text}"
                
                metadata["total_pages"] = len(pdf.pages)
        
        except Exception as e:
            # 如果pdfplumber失败，尝试PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    
                    for page_num, page in enumerate(pdf_reader.pages, 1):
                        page_text = page.extract_text()
                        if page_text:
                            content += f"\n\n[第{page_num}页]\n{page_text}"
                    
                    metadata["total_pages"] = len(pdf_reader.pages)
            
            except Exception as e2:
                logging.error(f"PDF处理失败: {e2}")
                return None
        
        if not content.strip():
            logging.warning(f"PDF文件无法提取文本: {file_path}")
            return None
        
        # 生成文档ID
        doc_id = self._generate_doc_id(file_path, content)
        
        return Document(
            id=doc_id,
            title=file_path.stem,
            content=content.strip(),
            file_path=str(file_path),
            file_type="pdf",
            metadata=metadata
        )
    
    def _process_txt(self, file_path: Path) -> Optional[Document]:
        """处理TXT文件
        
        Args:
            file_path: TXT文件路径
            
        Returns:
            文档对象
        """
        try:
            # 尝试多种编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
            content = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                logging.error(f"无法读取文件: {file_path}")
                return None
            
            metadata = {
                "source": file_path.name,
                "file_path": str(file_path),
                "file_type": "txt",
                "file_size": file_path.stat().st_size
            }
            
            # 生成文档ID
            doc_id = self._generate_doc_id(file_path, content)
            
            return Document(
                id=doc_id,
                title=file_path.stem,
                content=content.strip(),
                file_path=str(file_path),
                file_type="txt",
                metadata=metadata
            )
        
        except Exception as e:
            logging.error(f"TXT文件处理失败: {e}")
            return None
    
    def _process_markdown(self, file_path: Path) -> Optional[Document]:
        """处理Markdown文件
        
        Args:
            file_path: Markdown文件路径
            
        Returns:
            文档对象
        """
        # Markdown文件本质上是文本文件
        document = self._process_txt(file_path)
        if document:
            document.metadata["file_type"] = "markdown"
        return document
    
    def _process_docx(self, file_path: Path) -> Optional[Document]:
        """处理DOCX文件
        
        Args:
            file_path: DOCX文件路径
            
        Returns:
            文档对象
        """
        if not DOCX_AVAILABLE:
            logging.warning("Word文档处理库未安装")
            return None
        
        try:
            doc = DocxDocument(file_path)
            
            # 提取段落文本
            content = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content += paragraph.text + "\n\n"
            
            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = "\t".join([cell.text for cell in row.cells])
                    content += row_text + "\n"
                content += "\n"
            
            if not content.strip():
                logging.warning(f"DOCX文件无法提取文本: {file_path}")
                return None
            
            metadata = {
                "source": file_path.name,
                "file_path": str(file_path),
                "file_type": "docx",
                "file_size": file_path.stat().st_size
            }
            
            # 生成文档ID
            doc_id = self._generate_doc_id(file_path, content)
            
            return Document(
                id=doc_id,
                title=file_path.stem,
                content=content.strip(),
                file_path=str(file_path),
                file_type="docx",
                metadata=metadata
            )
        
        except Exception as e:
            logging.error(f"DOCX文件处理失败: {e}")
            return None
    
    def _generate_doc_id(self, file_path: Path, content: str) -> str:
        """生成文档ID
        
        Args:
            file_path: 文件路径
            content: 文档内容
            
        Returns:
            文档ID
        """
        # 使用文件路径和内容的哈希值作为ID
        hash_input = f"{file_path}_{content[:1000]}"
        return hashlib.md5(hash_input.encode()).hexdigest()
    
    def chunk_documents(self, documents: List[Document]) -> List[DocumentChunk]:
        """将文档分块
        
        Args:
            documents: 文档列表
            
        Returns:
            文档块列表
        """
        chunks = []
        
        for document in documents:
            doc_chunks = self._chunk_document(document)
            chunks.extend(doc_chunks)
        
        return chunks
    
    def _chunk_document(self, document: Document) -> List[DocumentChunk]:
        """将单个文档分块
        
        Args:
            document: 文档对象
            
        Returns:
            文档块列表
        """
        chunks = []
        content = document.content
        
        # 按段落分割
        paragraphs = content.split('\n\n')
        
        current_chunk = ""
        current_size = 0
        chunk_index = 0
        
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue
            
            paragraph_size = len(paragraph)
            
            # 如果当前块加上新段落超过大小限制
            if current_size + paragraph_size > self.chunk_size and current_chunk:
                # 保存当前块
                chunk = self._create_chunk(
                    document, 
                    current_chunk.strip(), 
                    chunk_index
                )
                chunks.append(chunk)
                
                # 开始新块（保留重叠）
                if self.chunk_overlap > 0:
                    overlap_text = current_chunk[-self.chunk_overlap:]
                    current_chunk = overlap_text + "\n\n" + paragraph
                    current_size = len(current_chunk)
                else:
                    current_chunk = paragraph
                    current_size = paragraph_size
                
                chunk_index += 1
            else:
                # 添加到当前块
                if current_chunk:
                    current_chunk += "\n\n" + paragraph
                else:
                    current_chunk = paragraph
                current_size += paragraph_size
        
        # 保存最后一个块
        if current_chunk.strip():
            chunk = self._create_chunk(
                document, 
                current_chunk.strip(), 
                chunk_index
            )
            chunks.append(chunk)
        
        return chunks
    
    def _create_chunk(self, document: Document, content: str, chunk_index: int) -> DocumentChunk:
        """创建文档块
        
        Args:
            document: 原文档
            content: 块内容
            chunk_index: 块索引
            
        Returns:
            文档块
        """
        chunk_id = f"{document.id}_chunk_{chunk_index}"
        
        # 复制文档元数据并添加块信息
        metadata = document.metadata.copy()
        metadata.update({
            "document_id": document.id,
            "document_title": document.title,
            "chunk_index": chunk_index,
            "chunk_size": len(content)
        })
        
        return DocumentChunk(
            id=chunk_id,
            content=content,
            metadata=metadata
        )